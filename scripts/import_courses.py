#!/usr/bin/env python3
"""Generate the Clearing course dataset from the approved Word document."""

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


TITLE_FIXES = {
    "Bioveterinary Science with Inegrated Master’s": "Bioveterinary Science with Integrated Master's",
    "Computer Science with Artificial Intelligence)": "Computer Science with Artificial Intelligence",
    "Computer Science with Digital Forensics)": "Computer Science with Digital Forensics",
    "Computer Science with Software Engineering)": "Computer Science with Software Engineering",
}

UNDERGRADUATE_STATUS_OVERRIDES = {
    "Medicine": "Full",
    "Midwifery": "Waiting list",
    "Nursing (Adult)": "Waiting list",
    "Nursing (Children's)": "Waiting list",
    "Nursing (Mental Health)": "Limited vacancies",
    "Occupational Therapy": "Waiting list",
    "Paramedic Science with Integrated Master's": "Limited vacancies",
    "Pharmacy": "Full",
    "Physiotherapy": "Full",
    "Radiography (Diagnostic Imaging)": "Full",
    "Social Work": "Waiting list",
    "Speech and Language Therapy with Integrated Master's (pre-registration)": "Limited vacancies",
    "Veterinary Medicine and Surgery": "Full",
}

FOUNDATION_STATUS_OVERRIDES = {
    "Nursing (Children's) with Foundation Year": "Limited vacancies",
    "Pharmacy with Foundation Year": "Full",
    "Physiotherapy with Foundation Year": "Full",
    "Radiography (Diagnostic Imaging) with Foundation Year": "Waiting list",
    "Speech and Language Therapy with Integrated Master's (pre-registration) with Foundation Year": "Waiting list",
    "Veterinary Medicine and Surgery with Foundation Year": "Full",
}

FOUNDATION_OFFER_OVERRIDES = {
    "Physiotherapy with Foundation Year": "96 UCAS tariff points",
    "Radiography (Diagnostic Imaging) with Foundation Year": "96 UCAS tariff points",
    "Social Work with Foundation Year": "40 UCAS tariff points",
    "Speech and Language Therapy with Integrated Master's (pre-registration) with Foundation Year": "56 UCAS tariff points",
}

UNDERGRADUATE_LINK_FALLBACKS = {
    "Zoology C302 with Integrated Master’s": (
        "Zoology with Integrated Master's",
        "https://www.keele.ac.uk/study/undergraduate-2026/undergraduatecourses/zoologywithintegratedmasters/",
    ),
}


def clean(value):
    return re.sub(r"\s+", " ", value or "").strip()


def first_hyperlink(paragraph, document):
    hyperlinks = paragraph._p.xpath(".//w:hyperlink")
    if not hyperlinks:
        return None, None

    hyperlink = hyperlinks[0]
    title = clean("".join(node.text or "" for node in hyperlink.xpath(".//w:t")))
    relationship_id = hyperlink.get(qn("r:id"))
    relationship = document.part.rels.get(relationship_id) if relationship_id else None
    return title, relationship.target_ref if relationship else None


def status_from_text(value):
    upper = value.upper()
    if "FULL" in upper or "NO VACANCIES" in upper:
        return "Full"
    if "WAITING LIST" in upper:
        return "Waiting list"
    if "LIMITED VACANCIES" in upper:
        return "Limited vacancies"
    return "Vacancies"


def ucas_code(value):
    match = re.match(r"\s*([A-Z][A-Z0-9]{3})", value)
    return match.group(1) if match else ""


def award(value):
    match = re.search(
        r"\b(?:BA|BSc|BEng|LLB|MSci|MEng|MChem|MMath|MCOMP|MBChB|MPharm|BVetMS|FdSc)"
        r"(?:\s*\(Hons\))?",
        value,
        re.IGNORECASE,
    )
    return clean(match.group(0)) if match else ""


def offer(value):
    match = re.search(r"(\d+)\s*UCAS\s*(?:tariff\s*)?points", value, re.IGNORECASE)
    if not match:
        match = re.search(r"(\d+)\s*points", value, re.IGNORECASE)
    if match:
        return f"{match.group(1)} UCAS tariff points"
    if re.search(r"pass\s+HND", value, re.IGNORECASE):
        return "Pass HND"
    if re.search(r"2:2\s+Degree", value, re.IGNORECASE):
        return "2:2 degree"
    return "See entry requirements"


def undergraduate_course(paragraph, document):
    title, url = first_hyperlink(paragraph, document)
    if not title or not url:
        for prefix, fallback in UNDERGRADUATE_LINK_FALLBACKS.items():
            if clean(paragraph.text).startswith(prefix):
                title, url = fallback
                break
    if not title or not url:
        return None

    title = TITLE_FIXES.get(title, title)
    full_text = clean(paragraph.text)
    remainder = clean(full_text.replace(title, "", 1))
    course_status = UNDERGRADUATE_STATUS_OVERRIDES.get(title, status_from_text(full_text))
    has_specific_requirements = "*" in full_text
    course_award = award(remainder)

    return {
        "title": title,
        "type": "Undergraduate",
        "status": course_status,
        "ucas": ucas_code(remainder),
        "requirements": offer(remainder),
        "summary": course_award or ("Top-up degree" if "Top-up" in title else "Undergraduate course"),
        "info": (
            "This course has additional subject-specific requirements. View the full course details "
            "or contact Admissions before applying."
            if has_specific_requirements
            else "Most undergraduate courses require GCSE English Language at grade 4/C or equivalent. "
            "Some courses also require GCSE Mathematics at grade 4/C or equivalent."
        ),
        "url": url,
    }


def foundation_course(paragraph, document):
    title, url = first_hyperlink(paragraph, document)
    if not title or not url:
        return None

    title = TITLE_FIXES.get(title, title)
    clean_url = url.split("#", 1)[0].rstrip("/") + "/#foundationyear"
    return {
        "title": title,
        "type": "Foundation year",
        "status": FOUNDATION_STATUS_OVERRIDES.get(title, "Vacancies"),
        "ucas": "",
        "requirements": FOUNDATION_OFFER_OVERRIDES.get(title, "See entry requirements"),
        "summary": "Foundation Year route",
        "info": (
            "Foundation Year requirements vary by course. View the full course details or call "
            "01782 491120 to discuss your qualifications."
        ),
        "url": clean_url,
    }


def extract_courses(source):
    document = Document(source)
    courses = []

    for paragraph in document.paragraphs[17:205]:
        course = undergraduate_course(paragraph, document)
        if course:
            courses.append(course)

    for paragraph in document.paragraphs[207:391]:
        course = foundation_course(paragraph, document)
        if course:
            courses.append(course)

    courses.sort(key=lambda item: (item["title"].casefold(), item["type"] != "Undergraduate"))
    return courses


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("--output", type=Path, default=Path("courses.js"))
    args = parser.parse_args()

    courses = extract_courses(args.source)
    payload = json.dumps(courses, ensure_ascii=False, indent=2)
    args.output.write_text(
        "// Generated from the approved Clearing 2026 Word document.\n"
        f"window.CLEARING_COURSES = Object.freeze({payload});\n",
        encoding="utf-8",
    )
    print(f"Wrote {len(courses)} courses to {args.output}")


if __name__ == "__main__":
    main()
